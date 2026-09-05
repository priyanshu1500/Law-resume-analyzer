"""Generate synthetic test-fixture resumes (PDF + DOCX) for the extraction/
evidence/rules pipeline. Not real people's data — entirely fabricated for
testing. Run: python gen_fixtures.py
"""
import os
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------- GOOD (single column)
def make_good_pdf():
    path = os.path.join(HERE, "good-single-column.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    y = h - 60

    def line(text, size=10, font="Helvetica", dy=16):
        nonlocal y
        c.setFont(font, size)
        c.drawString(60, y, text)
        y -= dy

    def header(text):
        nonlocal y
        y -= 6
        line(text, size=12, font="Helvetica-Bold", dy=18)

    line("Aditi Sharma", size=16, font="Helvetica-Bold", dy=20)
    line("aditi.sharma@nlu.ac.in  |  +91 98765 43210  |  New Delhi")

    header("EDUCATION")
    line("National Law University, Delhi — B.A., LL.B. (Hons.)  2022-2027")
    line("CGPA: 8.2 / 10")

    header("EXPERIENCE")
    line("Cyril Amarchand Mangaldas — Summer Intern, Corporate M&A  2026")
    line("- Drafted share purchase agreement conditions-precedent checklist for a Rs 400 crore deal")
    line("- Conducted due diligence across 60+ contracts and flagged 12 material risks")
    line("- Reviewed disclosure schedules and prepared a redline summary for the deal team")
    line("AZB & Partners — Winter Intern, Dispute Resolution  2025")
    line("- Researched arbitration precedent on investor-state disputes for a live matter")
    line("- Prepared a 15-page memo on limitation issues under the Arbitration Act")

    header("MOOT COURT & PUBLICATIONS")
    line("Runner-up, National Arbitration Moot Court Competition  2026")
    line("Publication: \"Conditions Precedent in Indian M&A Practice\", NLU Law Review (forthcoming)")

    header("SKILLS")
    line("Contract drafting, due diligence, legal research, Hindi and English")

    c.showPage()
    c.save()
    print("wrote", path)


# ---------------------------------------------------------------- BAD (two column, weak content)
def make_bad_pdf():
    path = os.path.join(HERE, "bad-two-column.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    left_x, right_x = 50, 320

    c.setFont("Helvetica-Bold", 14)
    c.drawString(left_x, h - 50, "Rohan Mehta")
    c.setFont("Helvetica", 9)
    c.drawString(left_x, h - 66, "Date of Birth: 14/03/2003   Gender: Male   Religion: Hindu")
    c.drawString(left_x, h - 80, "House No. 42, Sector 9, Rohini, New Delhi, 110085")
    c.drawString(left_x, h - 94, "rohan.mehta.law@gmail.com")

    # LEFT COLUMN — Experience (weak phrasing)
    y = h - 130
    c.setFont("Helvetica-Bold", 11); c.drawString(left_x, y, "EXPERIENCE"); y -= 16
    c.setFont("Helvetica", 9)
    for t in [
        "Some Law Firm — Intern (2025)",
        "- Responsible for legal research",
        "- Helped with various tasks",
        "- Worked on misc work assigned by seniors",
        "Another Firm — Intern (2024)",
        "- Assisted with legal research",
        "- Involved in general research",
    ]:
        c.drawString(left_x, y, t); y -= 14

    # RIGHT COLUMN — Education / Skills, interleaved by row position
    y2 = h - 130
    c.setFont("Helvetica-Bold", 11); c.drawString(right_x, y2, "EDUCATION"); y2 -= 16
    c.setFont("Helvetica", 9)
    for t in [
        "Some Law College (2021-2026)",
        "SKILLS",
        "MS Office, Westlaw, LexisNexis",
        "REFERENCES",
        "References available upon request",
    ]:
        c.drawString(right_x, y2, t); y2 -= 14

    c.showPage()
    c.save()
    print("wrote", path)


def make_good_docx():
    path = os.path.join(HERE, "good-single-column.docx")
    d = Document()
    d.add_heading("Aditi Sharma", level=1)
    d.add_paragraph("aditi.sharma@nlu.ac.in | +91 98765 43210 | New Delhi")
    d.add_heading("EDUCATION", level=2)
    d.add_paragraph("National Law University, Delhi — B.A., LL.B. (Hons.) 2022-2027")
    d.add_heading("EXPERIENCE", level=2)
    d.add_paragraph("Cyril Amarchand Mangaldas — Summer Intern, Corporate M&A 2026")
    d.add_paragraph("Drafted share purchase agreement conditions-precedent checklist for a Rs 400 crore deal", style="List Bullet")
    d.add_paragraph("Conducted due diligence across 60+ contracts and flagged 12 material risks", style="List Bullet")
    d.add_heading("SKILLS", level=2)
    d.add_paragraph("Contract drafting, due diligence, legal research")
    d.save(path)
    print("wrote", path)


if __name__ == "__main__":
    make_good_pdf()
    make_bad_pdf()
    make_good_docx()
